#!/usr/bin/env python
"""Build the evidence-backed SKRU-1 special section as a GOST-aligned DOCX."""

from __future__ import annotations

from datetime import datetime
from hashlib import sha256
import json
from pathlib import Path
from typing import Iterable, Sequence
from uuid import uuid4

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "thesis"
OUTPUT_PATH = OUTPUT_DIR / "SPECIAL_SECTION_SKRU1_RU.docx"
SOURCE_MAP_PATH = OUTPUT_DIR / "SPECIAL_SECTION_SKRU1_RU_SOURCE_MAP.json"


DATA_PATHS = {
    "gate_a1": ROOT / "artifacts" / "data_quality" / "gate_a1_report.json",
    "gate_b5": ROOT / "artifacts" / "model_selection" / "t1_b5_evidence_v1" / "gate_b5_report.json",
    "gate_b6": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "gate_b6_report.json",
    "b6_analytics": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "analytics_summary.json",
    "b6_temporal_metrics": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "temporal_aggregate_metrics.csv",
    "gate_c0": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "gate_c0_report.json",
    "gate_c0_validation": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "validation_report.json",
    "gate_c_architecture": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "architecture_eligibility.csv",
    "gate_c_lengths": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "sequence_length_distribution.csv",
    "gate_c1_validation": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "validation_report.json",
    "gate_c1_protocol_freeze": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "protocol_freeze.json",
    "gate_c1_admission": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "c2_admission_manifest.json",
    "gate_c1_label_access": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "outer_label_access_ledger.json",
    "gate_c1_registry": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "model_registry.json",
    "gate_c1_temporal_metrics": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "temporal_aggregate_metrics.csv",
    "gate_c1_fold_metrics": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "temporal_fold_metrics.csv",
    "gate_c1_seed_stability": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "seed_stability_metrics.csv",
    "gate_c1_native_metrics": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "student_t_native_metrics.csv",
    "gate_c1_compute": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "compute_resource_inventory.csv",
    "gate_c1_checkpoints": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "checkpoint_inventory.csv",
    "gate_c1_execution_incident": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "execution_incident_register.json",
    "gate_c1_screening": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "screening_register.csv",
    "gate_c1_worker_status": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "worker_status.csv",
    "gate_c1_tuning_inventory": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "tuning_inventory.csv",
    "gate_c1_scored_predictions": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "scored_temporal_predictions.csv",
    "gate_c1_artifact_inventory": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "artifact_inventory.csv",
    "gate_c1_figure_manifest": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "figure_manifest.json",
    "gate_c1_notebook_report": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "notebook_execution_report.json",
    "gate_c1_visual_qa": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "visual_qa_report.json",
    "gate_c1_environment": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "environment" / "hardware_report.json",
    "gate_c1_reader_manifest": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "reader_materials_manifest.json",
    "gate_c1_reporting_inventory": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "reporting_artifact_inventory.csv",
    "suite_v4": ROOT / "artifacts" / "governance" / "final_candidate_suite_v4.json",
    "holdout_v3": ROOT / "artifacts" / "governance" / "final_holdout_v3_status.json",
    "feature_contract": ROOT / "SKRU1_ACTUAL_DATA_TABLES_v1" / "02_eda_targets_v1" / "target_tables" / "formal_feature_contract.csv",
    "target_contract": ROOT / "SKRU1_ACTUAL_DATA_TABLES_v1" / "02_eda_targets_v1" / "target_tables" / "target_contract.json",
    "filatova_thesis": ROOT / "inputs" / "sources" / "primary" / "ВКР_Филатова_М_С.docx",
}


FIGURES = {
    "network": ROOT / "SKRU1_ACTUAL_DATA_TABLES_v1" / "01_reconstruction_v3_2" / "figures" / "02_survey_network_v3.png",
    "coverage": ROOT / "SKRU1_ACTUAL_DATA_TABLES_v1" / "01_reconstruction_v3_2" / "figures_v3_2" / "03_campaign_coverage.png",
    "b6_temporal": ROOT / "docs" / "thesis" / "figures" / "01_temporal_screen_mae.png",
    "b6_rolling": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "figures" / "02_rolling_mae_by_date.png",
    "b6_spatial": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "figures" / "03_spatial_stability.png",
    "b6_transition": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "figures" / "04_transition_error_heatmap.png",
    "b6_calibration": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "figures" / "05_conformal_calibration.png",
    "b6_learning": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "figures" / "06_learning_curves.png",
    "b6_sensitivity": ROOT / "artifacts" / "model_selection" / "t1_b6_expanded_v1" / "figures" / "07_profile_cluster_sensitivity.png",
    "c_length": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "figures" / "01_sequence_length_distribution.png",
    "c_gaps": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "figures" / "02_gap_and_missingness_geometry.png",
    "c_arch": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "figures" / "03_architecture_eligibility.png",
    "c_folds": ROOT / "artifacts" / "model_selection" / "t1_gate_c0_sequence_audit_v1" / "figures" / "04_fold_design.png",
    "c1_temporal": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "figures" / "01_ensemble_temporal_mae.png",
    "c1_rolling": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "figures" / "02_rolling_mae_by_target_date.png",
    "c1_seed": ROOT / "artifacts" / "model_selection" / "t1_gate_c1_compact_screen_v1" / "figures" / "03_seed_stability.png",
}


REFERENCES = [
    "ГОСТ 7.32–2017. Система стандартов по информации, библиотечному и издательскому делу. Отчет о научно-исследовательской работе. Структура и правила оформления. — Введ. 2018-07-01. — Москва : Стандартинформ, 2017.",
    "ГОСТ Р 7.0.100–2018. Система стандартов по информации, библиотечному и издательскому делу. Библиографическая запись. Библиографическое описание. Общие требования и правила составления. — Введ. 2019-07-01. — Москва : Стандартинформ, 2018.",
    "ГОСТ Р 7.0.5–2008. Система стандартов по информации, библиотечному и издательскому делу. Библиографическая ссылка. Общие требования и правила составления. — Введ. 2009-01-01. — Москва : Стандартинформ, 2008.",
    "SKRU-1 Data Foundation v3.2.1. Карточка набора данных : электронный ресурс. — Локальный репозиторий проекта, 2026. — Путь: SKRU1_ACTUAL_DATA_TABLES_v1/01_reconstruction_v3_2/DATASET_CARD.md.",
    "Методология реконструкции SKRU-1 v3.2 : электронный ресурс. — Локальный репозиторий проекта, 2026. — Путь: SKRU1_ACTUAL_DATA_TABLES_v1/01_reconstruction_v3_2/METHODOLOGY_V3_2.md.",
    "Филатова М. С. Геоинформационное моделирование и прогнозирование полей оседаний земной поверхности : выпускная квалификационная работа. — Пермь, 2026. — Локальная копия первичного источника проекта.",
    "Gate A1: качество данных и leakage-аудит T1/T5 : машинный отчет. — Локальный репозиторий проекта, 2026. — Путь: artifacts/data_quality/gate_a1_report.json.",
    "Gate B5: Evidence & Benchmark Protocol : машинный отчет. — Локальный репозиторий проекта, 2026. — Путь: artifacts/model_selection/t1_b5_evidence_v1/gate_b5_report.json.",
    "Gate B6: Expanded Classical, Probabilistic & Small-Data Screening : машинный отчет. — Локальный репозиторий проекта, 2026. — Путь: artifacts/model_selection/t1_b6_expanded_v1/gate_b6_report.json.",
    "Gate C: протокол sequence-моделей T1 : электронный ресурс. — Локальный репозиторий проекта, 2026. — Путь: docs/governance/GATE_C_PROTOCOL.md.",
    "Kalman R. E. A New Approach to Linear Filtering and Prediction Problems // Journal of Basic Engineering. — 1960. — Vol. 82, no. 1. — P. 35–45. — DOI 10.1115/1.3662552.",
    "Blom H. A. P., Bar-Shalom Y. The Interacting Multiple Model Algorithm for Systems with Markovian Switching Coefficients // IEEE Transactions on Automatic Control. — 1988. — Vol. 33, no. 8. — P. 780–783.",
    "Hyndman R. J., Koehler A. B. Another Look at Measures of Forecast Accuracy // International Journal of Forecasting. — 2006. — Vol. 22, no. 4. — P. 679–688. — DOI 10.1016/j.ijforecast.2006.03.001.",
    "Gneiting T., Raftery A. E. Strictly Proper Scoring Rules, Prediction, and Estimation // Journal of the American Statistical Association. — 2007. — Vol. 102, no. 477. — P. 359–378. — DOI 10.1198/016214506000001437.",
    "Cameron A. C., Miller D. L. A Practitioner’s Guide to Cluster-Robust Inference // Journal of Human Resources. — 2015. — Vol. 50, no. 2. — P. 317–372.",
    "Hochreiter S., Schmidhuber J. Long Short-Term Memory // Neural Computation. — 1997. — Vol. 9, no. 8. — P. 1735–1780. — DOI 10.1162/neco.1997.9.8.1735.",
    "Cho K. [et al.]. Learning Phrase Representations using RNN Encoder–Decoder for Statistical Machine Translation : arXiv:1406.1078. — 2014.",
    "Bai S., Kolter J. Z., Koltun V. An Empirical Evaluation of Generic Convolutional and Recurrent Networks for Sequence Modeling : arXiv:1803.01271. — 2018.",
    "Lim B., Arık S. Ö., Loeff N., Pfister T. Temporal Fusion Transformers for Interpretable Multi-horizon Time Series Forecasting // International Journal of Forecasting. — 2021. — Vol. 37, no. 4. — P. 1748–1764. — DOI 10.1016/j.ijforecast.2021.03.012.",
    "Gate C1: пятиseedовый compact sequence temporal screen : машинный отчет. — Локальный репозиторий проекта, 2026. — Путь: docs/reports/GATE_C1_COMPACT_SEQUENCE_SCREEN_RU.md.",
]


class SpecialSectionBuilder:
    def __init__(self) -> None:
        for path in [*DATA_PATHS.values(), *FIGURES.values()]:
            if not path.is_file():
                raise FileNotFoundError(path)
        self.a1 = json.loads(DATA_PATHS["gate_a1"].read_text(encoding="utf-8"))
        self.b5 = json.loads(DATA_PATHS["gate_b5"].read_text(encoding="utf-8"))
        self.b6 = json.loads(DATA_PATHS["gate_b6"].read_text(encoding="utf-8"))
        self.b6a = json.loads(DATA_PATHS["b6_analytics"].read_text(encoding="utf-8"))
        self.c0 = json.loads(DATA_PATHS["gate_c0"].read_text(encoding="utf-8"))
        self.c0_validation = json.loads(DATA_PATHS["gate_c0_validation"].read_text(encoding="utf-8"))
        self.architecture = pd.read_csv(DATA_PATHS["gate_c_architecture"])
        self.lengths = pd.read_csv(DATA_PATHS["gate_c_lengths"])
        self.c1_validation = json.loads(DATA_PATHS["gate_c1_validation"].read_text(encoding="utf-8"))
        self.c1_admission = json.loads(DATA_PATHS["gate_c1_admission"].read_text(encoding="utf-8"))
        self.c1_registry = json.loads(DATA_PATHS["gate_c1_registry"].read_text(encoding="utf-8"))
        self.c1_temporal = pd.read_csv(DATA_PATHS["gate_c1_temporal_metrics"])
        self.c1_folds = pd.read_csv(DATA_PATHS["gate_c1_fold_metrics"])
        self.c1_seeds = pd.read_csv(DATA_PATHS["gate_c1_seed_stability"])
        self.c1_native = pd.read_csv(DATA_PATHS["gate_c1_native_metrics"])
        self.c1_compute = pd.read_csv(DATA_PATHS["gate_c1_compute"])
        self.c1_checkpoints = pd.read_csv(DATA_PATHS["gate_c1_checkpoints"])
        self.c1_execution_incident = json.loads(
            DATA_PATHS["gate_c1_execution_incident"].read_text(encoding="utf-8")
        )
        self.c1_screen = pd.read_csv(DATA_PATHS["gate_c1_screening"])
        self.c1_workers = pd.read_csv(DATA_PATHS["gate_c1_worker_status"])
        self.c1_environment = json.loads(DATA_PATHS["gate_c1_environment"].read_text(encoding="utf-8"))
        deep = self.c1_temporal.loc[
            self.c1_temporal["model_id"].astype(str).str.startswith("C0")
            & self.c1_temporal["aggregation"].eq("mean_of_fixed_seeds")
        ]
        comparators = self.c1_temporal.loc[
            self.c1_temporal["model_id"].isin(
                ["B1_persistence_last_rate", "B7_two_regime_imm", "B8_student_t_robust_imm"]
            )
        ]
        self.c1_canonical = pd.concat((comparators, deep), ignore_index=True).sort_values(
            ["mae", "model_id"], kind="mergesort"
        )
        if len(self.c1_canonical) != 7 or self.c1_canonical["model_id"].nunique() != 7:
            raise RuntimeError("Gate C1 canonical temporal evidence is incomplete")
        if (
            len(self.c1_checkpoints) != 3860
            or int(self.c1_checkpoints["role"].eq("inner").sum()) != 3640
            or int(self.c1_checkpoints["role"].eq("outer").sum()) != 220
            or not self.c1_checkpoints["keep_top_k"].eq(5).all()
            or self.c1_checkpoints["outer_labels_used_for_ranking"].astype(bool).any()
        ):
            raise RuntimeError("Gate C1 checkpoint evidence is incomplete or unsafe")
        self.suite_v4 = json.loads(DATA_PATHS["suite_v4"].read_text(encoding="utf-8"))
        self.holdout_v3 = json.loads(DATA_PATHS["holdout_v3"].read_text(encoding="utf-8"))
        self.doc = Document()
        self.figure_number = 0
        self.table_number = 0
        self._configure_document()

    def _configure_document(self) -> None:
        document = self.doc
        section = document.sections[0]
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.header_distance = Mm(10)
        section.footer_distance = Mm(12)

        styles = document.styles
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(14)
        normal.font.color.rgb = RGBColor(0, 0, 0)
        _set_east_asia_font(normal, "Times New Roman")
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.first_line_indent = Cm(1.25)
        normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(0)
        normal.paragraph_format.widow_control = True

        heading1 = styles["Heading 1"]
        _style_heading(heading1, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, page_break=True, before=0, after=18)
        heading1.font.all_caps = True
        heading2 = styles["Heading 2"]
        _style_heading(heading2, size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT, page_break=True, before=0, after=12)
        heading3 = styles["Heading 3"]
        _style_heading(heading3, size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT, page_break=False, before=12, after=6)
        heading4 = styles["Heading 4"]
        _style_heading(heading4, size=14, alignment=WD_ALIGN_PARAGRAPH.LEFT, page_break=False, before=9, after=3)

        for name in ("List Bullet", "List Number"):
            style = styles[name]
            style.font.name = "Times New Roman"
            style.font.size = Pt(14)
            _set_east_asia_font(style, "Times New Roman")
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.left_indent = Cm(1.25)

        if "GOST Caption" not in styles:
            caption = styles.add_style("GOST Caption", WD_STYLE_TYPE.PARAGRAPH)
        else:
            caption = styles["GOST Caption"]
        caption.font.name = "Times New Roman"
        caption.font.size = Pt(12)
        _set_east_asia_font(caption, "Times New Roman")
        caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.first_line_indent = Cm(0)
        caption.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        caption.paragraph_format.space_before = Pt(3)
        caption.paragraph_format.space_after = Pt(9)
        caption.paragraph_format.keep_together = True

        if "GOST Table Text" not in styles:
            table_style = styles.add_style("GOST Table Text", WD_STYLE_TYPE.PARAGRAPH)
        else:
            table_style = styles["GOST Table Text"]
        table_style.font.name = "Times New Roman"
        table_style.font.size = Pt(12)
        _set_east_asia_font(table_style, "Times New Roman")
        table_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        table_style.paragraph_format.first_line_indent = Cm(0)
        table_style.paragraph_format.space_before = Pt(0)
        table_style.paragraph_format.space_after = Pt(0)

        properties = document.core_properties
        properties.title = "Специальная часть ВКР: прогнозирование скорости оседания SKRU-1"
        properties.subject = "Gate A/B evidence and Gate C0/C1 sequence research"
        properties.author = "Проект SKRU-1"
        properties.keywords = "SKRU-1; оседание; прогнозирование; Kalman; IMM; sequence models"
        properties.created = datetime(2026, 9, 1)
        properties.modified = datetime(2026, 9, 5)

        settings = document.settings._element
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        run = paragraph.add_run()
        _add_field(run, "PAGE")
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    def heading(self, text: str, level: int, *, page_break_before: bool | None = None) -> None:
        paragraph = self.doc.add_heading(text, level=level)
        paragraph.paragraph_format.left_indent = Cm(0)
        paragraph.paragraph_format.right_indent = Cm(0)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.keep_with_next = True
        if page_break_before is not None:
            paragraph.paragraph_format.page_break_before = page_break_before

    def paragraph(self, text: str, *, no_indent: bool = False, keep_with_next: bool = False) -> None:
        paragraph = self.doc.add_paragraph(text)
        if no_indent:
            paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.keep_with_next = keep_with_next

    def paragraphs(self, items: Iterable[str]) -> None:
        for item in items:
            self.paragraph(item)

    def bullets(self, items: Iterable[str]) -> None:
        for item in items:
            paragraph = self.doc.add_paragraph(item, style="List Bullet")
            paragraph.paragraph_format.keep_together = True

    def numbered(self, items: Iterable[str]) -> None:
        for index, item in enumerate(items, start=1):
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.left_indent = Cm(1.25)
            paragraph.paragraph_format.first_line_indent = Cm(-0.75)
            paragraph.paragraph_format.keep_together = True
            paragraph.add_run(f"{index}. {item}")

    def equation(self, expression: str, number: int | None = None) -> None:
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        paragraph.paragraph_format.keep_together = True
        text = expression if number is None else f"{expression}                                      ({number})"
        run = paragraph.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(14)
        run.italic = True

    def status_box(self, text: str) -> None:
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cell = table.cell(0, 0)
        cell.width = Cm(16.0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        _set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
        paragraph = cell.paragraphs[0]
        paragraph.style = self.doc.styles["GOST Table Text"]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.add_run(text)
        _set_cell_shading(cell, "F2F2F2")
        # The one-row status callout is announced as a labelled summary block
        # by assistive technologies rather than as an anonymous layout table.
        _set_repeat_table_header(table.rows[0])
        self.doc.add_paragraph().paragraph_format.space_after = Pt(0)

    def table(
        self,
        caption: str,
        headers: Sequence[str],
        rows: Sequence[Sequence[object]],
        widths_cm: Sequence[float] | None = None,
        *,
        keep_entire: bool = False,
    ) -> int:
        self.table_number += 1
        caption_paragraph = self.doc.add_paragraph(
            f"Таблица {self.table_number} — {caption}", style="GOST Caption"
        )
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        caption_paragraph.paragraph_format.keep_with_next = True
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_fixed_layout(table)
        if widths_cm is None:
            widths_cm = [16.0 / len(headers)] * len(headers)
        for index, header in enumerate(headers):
            cell = table.rows[0].cells[index]
            cell.width = Cm(widths_cm[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_text(cell, str(header), self.doc, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        _set_repeat_table_header(table.rows[0])
        _prevent_row_split(table.rows[0])
        for values in rows:
            row = table.add_row()
            _prevent_row_split(row)
            for index, value in enumerate(values):
                cell = row.cells[index]
                cell.width = Cm(widths_cm[index])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_text(cell, str(value), self.doc, alignment=alignment)
        if keep_entire:
            for row in table.rows[:-1]:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.keep_with_next = True
        return self.table_number

    def figure(self, path: Path, caption: str, *, width_cm: float = 15.8) -> int:
        self.figure_number += 1
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.keep_with_next = True
        inline_shape = paragraph.add_run().add_picture(str(path), width=Cm(width_cm))
        # python-docx does not expose alt text publicly.  Set the DrawingML
        # non-visual properties directly so every generated figure remains
        # accessible and the fix survives a full document rebuild.
        inline_shape._inline.docPr.set("descr", caption)
        inline_shape._inline.docPr.set("title", f"Рисунок {self.figure_number}")
        caption_paragraph = self.doc.add_paragraph(
            f"Рисунок {self.figure_number} — {caption}", style="GOST Caption"
        )
        return self.figure_number

    def build(self) -> None:
        self._opening()
        self._section_61()
        self._section_62()
        self._section_63()
        self._section_64()
        self._section_65()
        self._section_66()
        self._section_67()
        self._section_68()
        self._conclusion()
        self._references()

    def _opening(self) -> None:
        self.heading(
            "6 СПЕЦИАЛЬНАЯ ЧАСТЬ. ПРОГНОЗИРОВАНИЕ СКОРОСТИ ОСЕДАНИЯ ЗЕМНОЙ ПОВЕРХНОСТИ ПО ДАННЫМ МОНИТОРИНГА SKRU-1",
            1,
        )
        self.status_box(
            f'Статус доказательной базы на 05.09.2026: Gate A1 и Gate B5/B6 завершены; suite v4 заморожена с B7 в роли primary; Gate C0 имеет статус PASS_PROTOCOL_FROZEN; Gate C1 завершён как {self.c1_validation["status"]}, к Gate C2 допущено моделей: {len(self.c1_admission["admitted_model_ids"])}. Новый future/external holdout отсутствует и имеет статус PENDING_DATA.'
        )
        self.paragraphs(
            [
                "Раздел представляет воспроизводимое исследование прогнозирования скорости оседания земной поверхности для следующей плановой нивелирной кампании. Текст собран из проверяемых machine artifacts репозитория: контрактов данных, frozen split manifests, внешних и внутренних аудитов, fold-level predictions и агрегированных метрик. Числовые результаты в настоящей редакции не переносились вручную из промежуточных вычислений: источником служат опубликованные JSON/CSV-артефакты Gate A–C1 [7–10, 20].",
                "Доказательная граница принципиальна. Исходный набор является реконструированным исследовательским пакетом с синтетическими и восстановленными компонентами, а не новым эксплуатационным массивом реальных наблюдений. Поэтому полученные оценки характеризуют внутреннее поведение методов на зафиксированной геометрии данных, но не доказывают готовность модели к промышленному применению. Окончательная проверка возможна только после получения заранее замороженного future/external holdout.",
                "Оформление выполнено в рабочем GOST-aligned профиле: формат A4, поля 30/15/20/20 мм, шрифт Times New Roman 14 pt, полуторный интервал и абзац 1,25 см в соответствии с ГОСТ 7.32–2017 [1]. Библиографические описания и ссылки ориентированы на ГОСТ Р 7.0.100–2018 и ГОСТ Р 7.0.5–2008 [2, 3]. Если методические указания конкретного вуза устанавливают иные требования, они имеют приоритет и должны быть применены перед включением раздела в итоговую ВКР.",
            ]
        )
        self.paragraph("Структура специальной части", no_indent=True, keep_with_next=True)
        self.numbered(
            [
                "Постановка задачи, объект, предмет и ограничения исследования.",
                "Источники данных, provenance, канонические таблицы и качество.",
                "Реконструкция мониторинговой сети и зависимость наблюдений.",
                "Разведочный анализ, пропуски, режимы и формирование target.",
                "Алгоритмы прогнозирования: baselines, Kalman/IMM, classical ML, ENFS, uncertainty и compact sequence-модели Gate C1.",
                "Программная реализация, leakage guards и воспроизводимость.",
                "Экспериментальное исследование temporal, spatial, transition и probabilistic качества.",
                "Error Atlas, ограничения мониторинга и границы применимости.",
            ]
        )

    def _section_61(self) -> None:
        self.heading("6.1 Постановка задачи прогнозирования", 2)
        self.heading("6.1.1 Объект, предмет и практическая постановка", 3)
        self.paragraphs(
            [
                "Объектом исследования является пространственно-временная система геодезического мониторинга земной поверхности на территории, связанной с рудником СКРУ-1 Верхнекамского месторождения калийно-магниевых солей. Предмет исследования — методы краткосрочного прогнозирования скорости вертикального смещения рабочей точки к следующей плановой кампании при нерегулярном календаре, неодинаковой длине истории, пропусках и повторных наблюдениях внутри профиля.",
                "Основная задача T1 сформулирована как регрессия next-planned rate. Для каждого forecast origin известны точка, последняя доступная эпоха current, следующая фактически плановая целевая эпоха target и только те характеристики, которые могли быть доступны в момент формирования прогноза. Целевой показатель измеряется в миллиметрах в год; положительное направление соответствует принятой в target contract конвенции оседания. Горизонт индивидуален и вычисляется по календарной разности дат.",
            ]
        )
        self.equation("yᵢ = [sᵢ(tᵢᵗᵃʳ) − sᵢ(tᵢᶜᵘʳ)] · 365,25 / Δtᵢ", 1)
        self.paragraphs(
            [
                "Здесь sᵢ(tᵢᶜᵘʳ) и sᵢ(tᵢᵗᵃʳ) — наблюдённые накопленные оседания текущей и целевой эпох, а Δtᵢ — фактический горизонт в днях. Такая нормировка позволяет сравнивать origins с различными интервалами, но одновременно усиливает чувствительность к погрешности разности высот при коротком горизонте. Поэтому каждому target сопоставлена оценка sigma_rate_mm_y и clipped training weight, рассчитанные без использования скрытой истины.",
                "Вспомогательная задача T5 описывает раннее предупреждение о наступлении события в фиксированном горизонте. Gate A1 показал лишь 17 complete positive labels, поэтому T5 оставлена exploratory и не входит в текущую процедуру выбора primary. Специальная часть сосредоточена на T1, где доступны непрерывные target values и более содержательная temporal geometry.",
            ]
        )
        self.heading("6.1.2 Исследовательские вопросы и критерии успеха", 3)
        self.bullets(
            [
                "превосходит ли модель наивный persistence baseline на строго одинаковых rolling origins;",
                "сохраняется ли улучшение при равном весе профилей и зон, а не только в pooled micro-оценке;",
                "улучшается ли ошибка на accelerating и volatile-or-gap transitions без ухудшения стабильных режимов;",
                "калиброваны ли интервалы 50, 80 и 95%, и насколько ширина интервала соответствует достигнутому coverage;",
                "устойчивы ли выводы к выбору seed, target date, профиля, зоны и объёма доступной истории;",
                "можно ли до появления нового holdout заморозить одного кандидата без повторного использования старого test.",
            ]
        )
        self.paragraphs(
            [
                "Headline metric — MAE в мм/год. Дополнительно рассчитываются median absolute error, RMSE, signed и absolute bias, P90/P95/max absolute error, precision-weighted MAE/RMSE, skill относительно B1, MASE при наличии допустимого train-only denominator и direction-of-change accuracy. MAPE и sMAPE исключены, поскольку target содержит отрицательные и близкие к нулю значения; процентная нормировка в такой ситуации может становиться неопределённой или вводить в заблуждение [13].",
                "Финальный критерий отличается от внутреннего screening. Модель может стать только train-only internal primary, если одновременно проходит temporal, transition, profile, zone, calibration и reproducibility guards. Утверждение о реальном качестве до нового holdout запрещено независимо от значения внутренней MAE.",
            ]
        )
        self.table(
            "Границы допустимых научных утверждений",
            ["Уровень", "Разрешённый вывод", "Запрещённый вывод"],
            [
                ("Gate A/B", "Сравнение методов внутри frozen development geometry", "Промышленная точность на реальных будущих данных"),
                ("Gate C0", "Пригодность sequence representation и протокола", "Качество моделей до C1"),
                ("Gate C1", "Temporal train-only screening compact sequence-моделей", "Spatial/transition/final quality"),
                ("Suite v5", "Заранее выбранный train-only candidate", "Смена primary после holdout"),
                ("Future/external holdout", "Однократная внешняя оценка frozen primary", "Повторный тюнинг по результату holdout"),
            ],
            [2.6, 6.7, 6.7],
            keep_entire=True,
        )

    def _section_62(self) -> None:
        self.heading("6.2 Источники данных и provenance", 2)
        self.heading("6.2.1 Канонический контур", 3)
        self.paragraphs(
            [
                "Канонический модельный контур отделён от исторических и evaluation-only таблиц. Входы T1 образуют next_planned_features.csv, next_planned_operational_targets.csv, formal_feature_contract.csv и target_contract.json. Старые next_cycle_features.csv и next_cycle_targets.csv зарегистрированы только как historical comparison и не возвращаются model-facing loader. SHA-256 каждого файла фиксируется в Gate A1 report [7].",
                "Реконструкция опирается на первичный дипломный материал Филатовой, реестр источников, опубликованные опорные значения и воспроизводимые генераторы пакета v3.2 [4–6]. При этом наличие provenance code не превращает реконструированные поля в новые реальные измерения. В отчёте различаются документально транскрибированные, расчётные, реконструированные и синтетические компоненты; private_generation и evaluation_only слои исключены из features.",
            ]
        )
        canonical = self.a1["canonical_inputs"]
        self.table(
            "Канонические входы Gate A1",
            ["Роль", "Файл", "SHA-256, первые 12 символов"],
            [
                ("Признаки", Path(canonical["features"]["path"]).name, canonical["features"]["sha256"][:12]),
                ("T1 targets", Path(canonical["operational_targets"]["path"]).name, canonical["operational_targets"]["sha256"][:12]),
                ("Allowlist", Path(canonical["feature_contract"]["path"]).name, canonical["feature_contract"]["sha256"][:12]),
                ("Target contract", Path(canonical["target_contract"]["path"]).name, canonical["target_contract"]["sha256"][:12]),
                ("T5 labels", Path(canonical["early_warning_labels"]["path"]).name, canonical["early_warning_labels"]["sha256"][:12]),
            ],
            [3.0, 8.2, 5.3],
        )
        self.heading("6.2.2 Разбиения и зависимый grain", 3, page_break_before=True)
        split_rows = []
        for item in self.a1["split_manifests"]:
            if item["task"] == "T1":
                split_rows.append(
                    (
                        item["split"],
                        item["rows"],
                        item["points"],
                        item["profiles"],
                        f'{item["target_date_min"]} — {item["target_date_max"]}',
                        f'{100 * item["missing_feature_fraction"]:.2f}%',
                    )
                )
        self.table(
            "Frozen T1 manifests",
            ["Split", "Origins", "Точки", "Профили", "Target dates", "Пропуски features"],
            split_rows,
            [2.0, 2.0, 2.0, 2.0, 5.3, 3.2],
            keep_entire=True,
        )
        self.paragraphs(
            [
                "Полный canonical origin table содержит 1 274 строки, но это не 1 274 независимых объекта. Origins сворачиваются в 98 повторяющихся траекторий и 14 профилей, а даты кампаний общие для нескольких профилей. Внутри точки 34 split-safe static fields повторяются. Следовательно, random split искусственно переносит идентичность траектории и статический контекст из train в validation; обычный row-wise bootstrap занижает неопределённость.",
                "Gate A1 обнаружил 58 sample_id, чей последний token отражает историческую next-available кампанию, отличную от canonical target_campaign_id. Ключ остаётся уникальным, но должен трактоваться как непрозрачный join identifier. Любой разбор sample_id как смыслового признака запрещён; исправление требует новой версии IDs и всех manifests.",
            ]
        )
        reconciliation = self.a1["membership_reconciliation"]
        self.heading("6.2.3 Согласование 18 membership inconsistencies", 3)
        self.paragraphs(
            [
                f'Аудит явно связал {reconciliation["membership_rows_without_adjustment"]} membership rows без adjusted leveling epoch с {reconciliation["affected_model_origins"]} модельными origins без label. Разность grain объяснена: {reconciliation["reference_rows_outside_model_universe"]} строки относятся к reference points вне model universe, а {reconciliation["work_rows_without_eligible_origin"]} work rows не имеют допустимого предыдущего origin. Оставшиеся строки порождают шесть исключённых origins. Таким образом, значения 18 и 6 не противоречат друг другу и воспроизводятся join по campaign_id + point_id.',
                "Шесть origins исключены из loss; исходная таблица не изменена. Проверяемый mapping сохранён в artifacts/data_quality/membership_inconsistency_mapping.csv.",
            ]
        )

    def _section_63(self) -> None:
        self.heading("6.3 Реконструкция мониторинговой сети", 2)
        self.heading("6.3.1 Пространственная структура", 3)
        self.paragraphs(
            [
                "Наблюдательная система организована по профилям и рабочим точкам. Координаты используются для визуализации и формирования split-only spatial proxy, но не входят в estimator matrix. В версии spatial_quadrants_v1 четыре зоны определены медианными отсечениями локальных координат; это воспроизводимый тест пространственного переноса, а не инженерно утверждённое районирование.",
                "Неравномерность зон существенна: GEO_NE и GEO_SW содержат по 41 рабочей точке, GEO_NW и GEO_SE — по 8. Поэтому macro zone MAE и worst-zone рассматриваются вместе; pooled micro metric без этого может быть доминирован крупными зонами.",
            ]
        )
        self.figure(FIGURES["network"], "Реконструированная схема нивелирной сети и профилей")
        self.heading("6.3.2 Кампании, доступность и нерегулярность", 3)
        self.paragraphs(
            [
                "Календарь включает полные и focused кампании. Точка может быть наблюдена, не выбрана в focused campaign или пропущена по погоде, доступу, состоянию прибора либо разрушению знака. Для target formation следующая плановая эпоха определяется по фактически следующему наблюдению точки, а forecast_horizon_days сохраняет реальную календарную разность.",
                "Focused dates 17.01.2023 и 25.07.2023 представлены только одним–двумя зонами и двумя профилями. Они участвуют в rolling-origin оценке, но исключены из пространственной cross-validation. Пространственный аудит выполняется на последних трёх полных кампаниях: 18.10.2022, 16.05.2023 и 07.11.2023.",
            ]
        )
        self.figure(FIGURES["coverage"], "Покрытие точек по нивелирным кампаниям")
        self.table(
            "Уровни зависимости и требуемая оценка",
            ["Уровень", "Источник зависимости", "Контроль"],
            [
                ("Время", "Повтор точки по кампаниям", "Expanding rolling origin"),
                ("Профиль", "Общие геологические и календарные условия", "Spatio-temporal leave-profile-out"),
                ("Зона", "Пространственный перенос", "Spatio-temporal leave-zone-out"),
                ("Точка", "Повтор static features", "Worst 10% points; cluster sensitivity"),
            ],
            [2.8, 6.6, 7.1],
        )

    def _section_64(self) -> None:
        self.heading("6.4 Разведочный анализ и формирование target", 2)
        self.heading("6.4.1 Пропуски и drift", 3)
        self.paragraphs(
            [
                "Средняя доля пропущенных allowed feature cells мала — около 3,2–3,4% по T1 splits, однако распределение пропусков неоднородно. terrain_TRI_relative отсутствует в 32,418% origins, а lithology uncertainty — во всех строках. Поэтому simple global complete-case analysis недопустим: он меняет пространственный состав и может выбрасывать целые feature families.",
                "Preprocessing применяет train-fitted imputation и missing indicators. Категориальные уровни также фиксируются только по train; неизвестная категория получает отдельный код. Ни imputer, ни scaler, ни encoder не видят validation/test labels. Наибольший train-to-test drift среди числовых полей отмечен для n_history, |SMD| = 2,320, что подтверждает необходимость temporal validation и осторожность при переносе на поздние кампании.",
            ]
        )
        self.heading("6.4.2 Режимы движения и transition proxy", 3)
        self.paragraphs(
            [
                "Transition labels для анализа не являются скрытой физической истиной. Они вычисляются только из origin-available характеристик: recent_acceleration_mm_y2, std_last_3_rates_mm_y и missing_campaigns_since_previous. Порог абсолютного ускорения и порог волатильности оцениваются внутри соответствующего train role по 80-му процентилю. Эта процедура исключает использование будущего target при сегментации.",
                "Origins делятся на stable, accelerating, decelerating и volatile-or-gap. Pooled transition объединяет три нестабильных сегмента. Сегмент считается пригодным для selection только при поддержке не менее 20 origins и 5 профилей; иначе результат получает метку DESCRIPTIVE_LOW_SUPPORT и не заменяется искусственным штрафом.",
            ]
        )
        self.heading("6.4.3 Target, uncertainty и веса", 3)
        self.paragraphs(
            [
                "Primary target observed_rate_mm_y основан на разности двух adjusted settlement values. Его неопределённость sigma_rate_mm_y получается распространением стандартных неопределённостей текущей и целевой эпох через разность и индивидуальный горизонт. Training weight использует обратную дисперсию с clipping, чтобы одна формально точная строка не доминировала в objective.",
                "Target distribution train содержит отрицательные и близкие к нулю скорости, а также правый хвост до 151,5 мм/год. Поэтому MAE выбрана headline metric, RMSE и P95 показывают тяжесть хвоста, а signed bias контролирует систематическое завышение или занижение. R² публикуется только descriptively и не является критерием выбора.",
            ]
        )
        train_split = next(item for item in self.a1["split_manifests"] if item["task"] == "T1" and item["split"] == "train")
        distribution = json.loads(train_split["target_distribution_json"])
        self.table(
            "Распределение T1 target в train",
            ["Показатель", "Значение, мм/год"],
            [
                ("Минимум", f'{distribution["min"]:.2f}'),
                ("P05", f'{distribution["p05"]:.2f}'),
                ("Медиана", f'{distribution["median"]:.2f}'),
                ("Среднее", f'{distribution["mean"]:.2f}'),
                ("P95", f'{distribution["p95"]:.2f}'),
                ("Максимум", f'{distribution["max"]:.2f}'),
            ],
            [8.0, 8.0],
        )

    def _section_65(self) -> None:
        self.heading("6.5 Методы прогнозирования", 2, page_break_before=True)
        self.heading("6.5.1 Persistence и статистические baselines", 3)
        self.paragraphs(
            [
                "B1 persistence_last_rate переносит последнюю наблюдённую скорость на следующий горизонт. Несмотря на простоту, это обязательная опора: временные ряды коротки, а локальная скорость содержит сильную инерцию. Любая сложная модель должна сравниваться с B1 на точно одинаковых origins, иначе skill неинтерпретируем.",
                "B3 profile_robust_trend смешивает последнюю скорость, среднее последних трёх скоростей и профильное среднее, добавляя ограниченную компоненту ускорения. Модель использует только origin-available aggregates и не запоминает profile_id. ETS, ARIMA/ARIMAX и profile VAR формально исключены в Gate B5: на точку приходится 3–16 наблюдений с интервалами 42–560 дней и пропущенными кампаниями; равномерная интерполяция не обоснована, а профильный VAR при 14 профилях и неполной синхронности неустойчив.",
            ]
        )
        self.heading("6.5.2 Фильтр Калмана и IMM", 3)
        self.paragraphs(
            [
                "B5 реализует fixed Kalman model в пространстве settlement–velocity. Переход зависит от фактического delta_t, measurement variance задаётся uncertainty adjusted observation, process variance заморожена до оценки. B6 adaptive Kalman изменяет process/measurement response по innovation diagnostics внутри train-only protocol. Теоретическая основа линейной фильтрации приведена в [11].",
                "B7 two-regime IMM расширяет состояние до settlement–velocity–acceleration и объединяет stable и transition filters через марковские вероятности режима [12]. Переходы являются статистическими состояниями фильтра, а не идентифицированными геомеханическими стадиями. B8 сохраняет динамику B7, но вводит Student-t robust innovation weight; train-only audit показал активацию downweighting, однако не дал требуемого улучшения volatile-or-gap, поэтому B8 остался context comparator.",
                "Сильная сторона IMM на этих данных — явная обработка нерегулярного времени и малая параметрическая сложность. Ограничение — чувствительность к заранее выбранной структуре состояний и невозможность автоматически извлекать нелинейные зависимости из полного набора static/context features.",
            ]
        )
        self.equation("xₖ = F(Δtₖ)xₖ₋₁ + wₖ;      zₖ = Hxₖ + vₖ", 2)
        self.equation("p(xₖ | z₁:ₖ) = Σⱼ μₖʲ pⱼ(xₖ | z₁:ₖ)", 3)
        self.table(
            "Неизменяемые динамические comparators",
            ["ID", "Семейство", "Назначение", "Статус"],
            [
                ("B1", "Persistence", "Наивная последняя скорость", "Comparator"),
                ("B5", "Fixed Kalman", "Линейное состояние и uncertainty", "Comparator"),
                ("B6", "Adaptive Kalman", "Innovation-adaptive response", "Comparator"),
                ("B7", "Two-regime IMM", "Stable/transition mixture", "Suite v4 primary"),
                ("B8", "Student-t robust IMM", "Robust observation channel", "Context only"),
            ],
            [1.8, 4.0, 7.2, 3.0],
        )
        self.heading("6.5.3 Classical ML, probabilistic models и ENFS", 3)
        self.paragraphs(
            [
                "Gate B6 выполнил широкое nested train-only сравнение linear, kernel, longitudinal, boosting, glassbox, probabilistic и малых neural/tabular моделей. SAFE_ALL использовал полный executable allowlist, DYNAMIC_CORE_17 — заранее определённый динамический subset, NATIVE_CATEGORICAL — train-fitted categorical schema. IDs и zone/campaign keys оставались metadata; point_id применялся в GEE только как working-correlation group и отсутствовал в exogenous matrix.",
                "В screening вошли ElasticNet, Huber, RBF-SVR, Gaussian Process, Gaussian GEE, HistGradientBoosting, XGBoost, LightGBM, CatBoost, EBM, NGBoost, residual MLP и protocol-safe ENFS replica. ENFS воспроизводит архитектурную идею предшествующей ВКР [6], но не заявляет воспроизведение её численных результатов.",
                "Гиперпараметры выбирались только по трём forward-only inner folds каждого outer train. Для probabilistic models использовался CRPS с guardrail: point MAE не хуже B1 более чем на 5%. Early stopping для boosters и neural methods не использовал outer labels.",
            ]
        )
        self.heading("6.5.4 Sequence representation, compact GRU и LSTM", 3)
        self.status_box(
            f'Gate C0: PASS_PROTOCOL_FROZEN; Gate C1: {self.c1_validation["status"]}. Все приведённые deep-model metrics относятся только к nested temporal screen внутри t1_v1/train.'
        )
        self.paragraphs(
            [
                f'Для {self.c0["data"]["origins"]} train origins построено {self.c0["data"]["normalized_rows"]} normalized sequence rows: 16 позиций на origin с left padding. Фактическая длина истории 3–16, медиана {self.c0["data"]["history_length_median"]:.0f}. Последний actual token совпадает с current_date; target observation и все observation после origin отсутствуют.',
                "Шесть model channels повторяют семантику полей formal_feature_contract: settlement, interval rate, standard uncertainty, delta_t, missing-campaign count и campaign type. Padding, observation и missing masks передаются отдельно. point/profile/zone/campaign IDs сохраняются только для доказательства split geometry и не могут попасть в network matrix.",
                f'Нерегулярность значительна: положительный delta_t изменяется от {self.c0["data"]["delta_t_days_min_positive"]} до {self.c0["data"]["delta_t_days_max"]} дней при медиане {self.c0["data"]["delta_t_days_median_positive"]:.0f}; максимальный gap равен {self.c0["data"]["missing_campaigns_max"]} кампаниям. Эта геометрия поддерживает masked irregular-time models и не оправдывает скрытую интерполяцию.',
            ]
        )
        self.figure(FIGURES["c_length"], "Распределение фактической длины sequence history")
        self.figure(FIGURES["c_gaps"], "Нерегулярность интервалов и missing-campaign gaps")
        c1_metrics = self.c1_canonical.set_index("model_id")
        c1_screen = self.c1_screen.set_index("model_id")
        c1_compute = self.c1_compute.set_index("model_id")
        c1_status_display = {
            "PASSED_TEMPORAL_SCREEN": "PASSED",
            "REJECTED_TEMPORAL_SCREEN": "REJECTED",
            "REJECTED_MODEL_EXECUTION": "EXECUTION REJECT",
        }
        compact_descriptions = {
            "C01_compact_gru": ("GRU, последний valid hidden", "Huber; inner MAE"),
            "C02_compact_lstm": ("LSTM, последний valid hidden", "Huber; inner MAE"),
            "C03_causal_tcn": ("2 causal residual Conv1d blocks", "Huber; inner MAE"),
            "C04_probabilistic_gru_student_t": ("GRU + loc/scale/df head", "Student-t NLL; inner CRPS"),
        }
        self.table(
            "Компактные sequence-архитектуры Gate C1",
            ["Model", "Encoder/head", "Objective/selection", "Grid", "Параметры", "C1 status"],
            [
                (
                    model_id.split("_")[0],
                    compact_descriptions[model_id][0],
                    compact_descriptions[model_id][1],
                    len(next(item for item in self.c1_registry["models"] if item["model_id"] == model_id)["parameter_grid"]),
                    f'{int(c1_compute.loc[model_id, "parameter_count_min"])}–{int(c1_compute.loc[model_id, "parameter_count_max"])}',
                    c1_status_display[str(c1_screen.loc[model_id, "status"])],
                )
                for model_id in compact_descriptions
            ],
            [1.5, 4.1, 3.8, 1.2, 2.6, 3.3],
        )
        self.paragraphs(
            [
                f'C01 compact GRU достигла canonical mean-of-five-seeds MAE {c1_metrics.loc["C01_compact_gru", "mae"]:.3f} мм/год при median fold MAE {c1_metrics.loc["C01_compact_gru", "median_fold_mae"]:.3f}; temporal status — {c1_screen.loc["C01_compact_gru", "status"]}.',
                f'C02 compact LSTM получила MAE {c1_metrics.loc["C02_compact_lstm", "mae"]:.3f} мм/год и median fold MAE {c1_metrics.loc["C02_compact_lstm", "median_fold_mae"]:.3f}; temporal status — {c1_screen.loc["C02_compact_lstm", "status"]}. Различия между GRU и LSTM на этой выборке нельзя трактовать как общий рейтинг архитектур: независимых temporal units только 11, а пространственный audit ещё не выполнен.',
                "Для recurrent adapters left-padded tokens перед packing сдвигаются в начало без изменения хронологического порядка. Числовые каналы и target масштабируются только по соответствующему train role; неизвестная campaign category получает заранее предусмотренный unknown bucket.",
            ]
        )
        self.heading("6.5.5 Causal TCN и probabilistic Student-t GRU", 3)
        self.paragraphs(
            [
                f'C03 causal TCN использует два residual causal Conv1d blocks с dilation 1/2 и повторным masking padded activations. Её canonical MAE составила {c1_metrics.loc["C03_causal_tcn", "mae"]:.3f} мм/год, temporal status — {c1_screen.loc["C03_causal_tcn", "status"]}. Causal left convolution исключает доступ к будущим token positions; изменение padding values проверяется отдельным инвариантным тестом.',
                f'C04 probabilistic GRU выдаёт параметры Student-t: loc без ограничения, scale = softplus(raw)+10⁻³ и df = 2,01+softplus(raw). Point prediction равен loc в мм/год; canonical point MAE составила {c1_metrics.loc["C04_probabilistic_gru_student_t", "mae"]:.3f}, temporal status — {c1_screen.loc["C04_probabilistic_gru_student_t", "status"]}. Native CRPS/NLL и coverage публикуются отдельно по каждому seed; параметры пяти распределений не усредняются в псевдо-Student-t.',
                "Пять seeds 42117–42121 участвуют как в inner tuning, так и в outer refit; ни один seed не выбирается отдельно. C01–C03 выбираются по pooled inner MAE, C04 — по pooled inner CRPS с MAE tie-breaker. Outer epoch count равен медиане 15 inner best epochs выбранной configuration.",
                "TSMixer и compact TFT остаются условными будущими гипотезами и не входят в C1. N-BEATS, N-HiTS, PatchTST и iTransformer сохраняют статус NOT_ELIGIBLE_DATA_GEOMETRY. Leave-profile, leave-zone, transition audit, conformal calibration и suite v5 явно отложены до Gate C2.",
            ]
        )
        self.figure(FIGURES["c_arch"], "Предварительная пригодность sequence-архитектур")
        self.figure(FIGURES["c_folds"], "Замороженная outer-fold geometry Gate C")
        status_labels = {
            "REQUIRED_COMPACT_SCREEN": "Обязательный screen",
            "CONDITIONAL_COMPACT_SCREEN": "Условный screen",
            "NOT_ELIGIBLE_DATA_GEOMETRY": "Не допущена",
        }
        architecture_notes = {
            "C01_compact_gru": "Безусловно включена",
            "C02_compact_lstm": "Безусловно включена",
            "C03_causal_tcn": "Безусловно включена",
            "C04_probabilistic_gru_student_t": "Безусловно включена",
            "C05_tsmixer_compact": "После C0 mask audit",
            "C06_tft_compact": "Не более 100 тыс. параметров; без ID embeddings",
            "C07_nbeats": "Требует недоказанной regular grid",
            "C08_nhits": "Downsampling неидентифицируем при 3–16 наблюдениях",
            "C09_patchtst": "Слишком короткий sparse patch context",
            "C10_itransformer": "Недостаточная sample/channel geometry",
        }
        arch_rows = []
        for row in self.architecture.itertuples(index=False):
            arch_rows.append(
                (
                    row.model_id,
                    status_labels[row.status],
                    architecture_notes[row.model_id],
                )
            )
        self.table(
            "Architecture registry Gate C0",
            ["Model ID", "Статус", "Условие или основание"],
            arch_rows,
            [4.2, 4.5, 7.8],
        )
        self.heading("6.5.6 Методы, исключённые до обучения", 3)
        self.paragraphs(
            [
                "Предварительное исключение архитектуры — часть научного протокола. Оно предотвращает механическое увеличение model zoo за счёт методов, предпосылки которых не согласованы с данными. Для ETS/ARIMA/VAR основание — короткие нерегулярные ряды; для patch/hierarchical transformer-like models — недостаточная длина контекста и отсутствие обоснованной регулярной сетки.",
                "Исключённая модель может быть возвращена только в новой версии протокола после появления более длинных последовательностей, domain-governed interpolation scheme или внешнего набора с достаточным числом независимых trajectories. Изменять t1_train_gate_c_v1 задним числом запрещено.",
            ]
        )
        self.heading("6.5.7 Интерпретируемость и диагностические outputs", 3)
        self.paragraphs(
            [
                "Интерпретируемость оценивается на уровне механизмов и ошибок, а не только feature importance. Для B7 доступны regime probability и entropy; для EBM — shape functions; для trees — permutation/SHAP-like diagnostics только на train-only folds; для sequence models планируются saliency по time tokens и mask ablations. Идентификаторы не допускаются даже ради улучшения fit, поскольку такое улучшение означало бы memorization.",
                "Любая локальная интерпретация сопровождается support metadata: profile, zone, target date, history length, uncertainty и gap bin. Категории с менее чем 20 origins или пятью профилями не используются в model selection.",
            ]
        )
        self.heading("6.5.8 Правило выбора и frozen suite", 3)
        self.paragraphs(
            [
                "Единого weighted score нет. Eligibility проверяется как набор одновременных ограничений, затем eligible models упорядочиваются лексикографически по rolling MAE, transition MAE, worst-zone MAE, 95% WIS, fit time и model ID; Gate C добавляет seed IQR и parameter count. Такой порядок прозрачен и воспроизводится из machine metrics.",
                "Gate B6 завершился PASS_NO_NEW_PRIMARY: ни одна новая модель не прошла все criteria, поэтому B7 остался primary suite v4. Gate C заранее сохраняет тот же fallback. Если deep-модель не проходит хотя бы один guard, отсутствие нового primary считается корректным научным результатом.",
            ]
        )
        self.heading("6.5.9 Неопределённость и интервальная калибровка", 3)
        self.paragraphs(
            [
                "Native probabilistic outputs и common conformalized intervals публикуются отдельно. Общий scaled conformal wrapper получает residuals только из inner rolling OOF predictions; scale зависит от train-only horizon и uncertainty и ограничивается диапазоном 0,25–4. Outer validation labels не участвуют в calibration.",
                "Оцениваются empirical coverage 50, 80 и 95%, mean/median width, central interval score, weighted interval score, CRPS и NLL только внутри сопоставимой distribution family. CRPS и interval score являются proper scoring rules: они совместно учитывают calibration и sharpness [14].",
            ]
        )
        self.equation("ISα(l, u; y) = (u − l) + 2/α·(l − y)·1{y<l} + 2/α·(y − u)·1{y>u}", 4)

    def _section_66(self) -> None:
        self.heading("6.6 Программная реализация и воспроизводимость", 2)
        self.heading("6.6.1 Архитектура репозитория", 3)
        self.paragraphs(
            [
                "Репозиторий разделён на configs, src/skru1, scripts, tests, artifacts, notebooks, docs, inputs, requirements и work. Код обнаруживает root по pyproject.toml и gate_a1.yaml; абсолютные host paths в конфигурациях запрещены. Canonical loader является единственной точкой model-facing доступа к таблицам.",
                "Каждый gate имеет фазы freeze, analyze, validate и all либо эквивалентные model phases. Промежуточные checkpoints создаются только в work, а финальные CSV/JSON/joblib/notebook artifacts публикуются атомарно. Для frozen manifests повторный запуск допустим только при байтово идентичном содержимом; любое отличие требует новой версии.",
            ]
        )
        self.table(
            "Основные программные интерфейсы",
            ["Интерфейс", "Назначение", "Leakage boundary"],
            [
                ("CanonicalBundle", "Канонические tables/contracts", "Нет private/evaluation-only fields"),
                ("ManifestDataset", "Загрузка только через frozen sample IDs", "Test требует candidate record"),
                ("BenchmarkPlan", "Outer/inner folds и expected hashes", "Random split запрещён"),
                ("FitContext", "Train hash, feature hash, seed, groups", "Groups отдельно от X"),
                ("PredictionBundle", "Единая prediction schema", "Exact sample hash и no duplicates"),
                ("SequenceBundle", "Causal padded histories", "No future/target observation"),
                ("SequenceModelSpec", "Architecture, grid, seeds, channels", "Canonical spec SHA-256"),
                ("SequenceTensorBatch", "Channels, masks, lengths", "IDs хранятся отдельно от tensor"),
                ("SequenceTargetScaler", "Train-only target mean/std", "Fit provenance и state hash"),
                ("SequencePredictionBundle", "Unlabeled deep predictions", "Worker shard не содержит y_true"),
                ("C1BenchmarkPlan", "11 outer и 33 logical inner contexts", "Forward-only sample/sequence hashes"),
                ("TemporalAdmissionRecord", "Программный C2 admission", "Только PASSED_TEMPORAL_SCREEN"),
            ],
            [3.6, 6.8, 6.1],
        )
        self.heading("6.6.2 Автоматические leakage guards", 3)
        self.bullets(
            [
                "unique sample_id и отсутствие overlap между split/fold roles;",
                "T1 split строго по target_date, T5 — по label_horizon_end;",
                "forecast_horizon_days > 0;",
                "отсутствие true_, hidden, private, generator, event_onset_date, process_family и regime_stage;",
                "точное соответствие executable feature allowlist;",
                "запрет point/profile/zone/campaign IDs в estimator matrix;",
                "fit preprocessing только на provenance=train;",
                "ordinary KFold/random split завершаются исключением;",
                "test loader sealed до frozen candidate;",
                "Gate C history заканчивается на current_date и не содержит target observation;",
                "early stopping разрешён только во внутреннем rolling fold outer train;",
                "outer-validation labels отсутствуют в worker shard и присоединяются независимым scorer только после hash freeze всех 44 shards;",
                "runtime network guard завершает C1 job при попытке socket или URL access.",
            ]
        )
        self.heading("6.6.3 Среды выполнения и hardware", 3)
        self.paragraphs(
            [
                "Классические и probabilistic models разделены на b6_cpu и b6_ngboost из-за несовместимых диапазонов scikit-learn. MLP/ENFS использовали b6_torch. Для C1 создана свежая отдельная среда gate_c_torch строго из неизменяемого lock; B6 torch environment не использовалась как execution authority. Boosters обучаются на CPU для детерминированности; CUDA разрешена sequence models.",
                f'Авторитетный C1 screen выполнен на {self.c1_environment["gpu_name"]} с {self.c1_environment["gpu_memory_mib"]:.0f} MiB VRAM, Python {self.c1_environment["python_version"]}, PyTorch {self.c1_environment["torch_version"]} и CUDA {self.c1_environment["torch_cuda_version"]}; driver {self.c1_environment["gpu_driver_version"]}. Зафиксированы pip freeze, wheel URLs/SHA-256, OS/CPU/RAM/GPU capture, serialization smoke и два последовательных deterministic CUDA fits с tolerance 10⁻⁶. Mixed precision и TF32 выключены, DataLoader workers = 0.',
            ]
        )
        self.table(
            "Замороженные среды исследования",
            ["Environment", "Основные семейства", "Устройство", "Примечание"],
            [
                ("b6_cpu", "Linear, GPR/GEE, trees, EBM", "CPU", "Deterministic single-thread where required"),
                ("b6_ngboost", "NGBoost", "CPU", "Отдельный sklearn-compatible lock"),
                ("b6_torch", "Residual MLP, ENFS", "CUDA/CPU", "Историческая B6 среда"),
                ("gate_c_torch", "GRU, LSTM, TCN, Student-t GRU", "RTX 5070 Ti", "PyTorch 2.13.0+cu130; deterministic"),
            ],
            [3.0, 5.2, 2.6, 5.7],
        )
        self.heading("6.6.4 Checkpoint/recovery и CUDA-оптимизация", 3)
        benchmark = self.c1_execution_incident["incidents"][-1]["matched_runtime_benchmark"]
        checkpoint_bytes = int(self.c1_checkpoints["ranked_checkpoint_bytes"].sum())
        self.paragraphs(
            [
                "Для каждого C1 fit заморожена восстанавливаемая checkpoint-схема: model, optimizer, shuffle generator, Python/NumPy/Torch CPU/CUDA RNG states и provenance. Recovery checkpoint атомарно записывается после каждой завершённой стадии в 50 эпох и на terminal epoch. Каждый manifest хранит SHA-256 и пять полных ranked states.",
                "Для inner fits top-5 ранжируются по заранее заданной early-stopping metric, а для prediction восстанавливается rank 1. В outer refit хранятся последние пять эпох, но выбор всегда остаётся на preregistered final epoch. Outer labels не участвуют ни в ранжировании, ни в выборе checkpoint.",
                f'Инвентаризация содержит 3 860 manifests: 3 640 inner и 220 outer, всего 19 300 retained states объёмом {checkpoint_bytes / (1024 ** 3):.2f} ГиБ. Бинарные checkpoints хранятся только в gitignored work/; в artifacts публикуются лишь их машинная инвентаризация.',
                f'Без изменения model logic recurrent input переведён на векторизованный device-side gather; убраны per-row CPU/CUDA synchronization, AdamW запускается в fused CUDA mode, validation metrics считаются на GPU. В matched benchmark на {int(benchmark["fits_each_run"])} одинаковых fits mean time снижено с {benchmark["old_mean_fit_seconds"]:.3f} до {benchmark["new_mean_fit_seconds_including_checkpointing"]:.3f} с ({benchmark["mean_speedup_ratio"]:.2f}×), median — с {benchmark["old_median_fit_seconds"]:.3f} до {benchmark["new_median_fit_seconds_including_checkpointing"]:.3f} с ({benchmark["median_speedup_ratio"]:.2f}×). Новые timings уже включают top-5 checkpoint I/O.',
                f'Полное насыщение 16 ГиБ VRAM для этой задачи не ожидается: peak tensor allocation равна {self.c1_compute["peak_vram_mb"].max():.1f} MB, крупнейшая configuration содержит {int(self.c1_compute["parameter_count_max"].max())} параметров, sequence length не превышает 16, batch size заморожен на 32 и разрешён один deterministic GPU worker. Искусственное увеличение batch или параллельный запуск folds изменили бы execution semantics либо ослабили воспроизводимость. Поэтому корректный результат оптимизации — измеренное сокращение времени, а не максимальная занятость памяти GPU.',
            ]
        )
        self.table(
            "Checkpoint inventory Gate C1",
            ["Role", "Manifests", "Retained states", "Ranking", "Selected state"],
            [
                ("inner", "3 640", "18 200", "Frozen validation metric", "Rank 1"),
                ("outer", "220", "1 100", "Последние пять epochs", "Fixed final epoch"),
            ],
            [2.0, 2.3, 3.0, 4.5, 4.5],
        )

    def _section_67(self) -> None:
        self.heading("6.7 Экспериментальное исследование", 2)
        self.heading("6.7.1 Протокол B5/B6 и Gate C1", 3)
        self.paragraphs(
            [
                "B5/B6 используют только 911 строк t1_v1/train. Broad temporal evidence образуют 11 rolling target dates с 18.05.2021 по 07.11.2023. Перед первым outer fold доступны 316 origins из восьми предыдущих target campaigns. Каждый outer train содержит последние три допустимые forward-only inner folds для tuning.",
                "Spatial design образуют 14 профилей × 3 полные кампании и 4 proxy-зоны × 3 кампании. Held group полностью исключается из outer train и всех inner folds. После shortlist строятся learning curves на audit tail 07.11.2023 с 217, 423, 708 и 823 train origins; гиперпараметры не перенастраиваются.",
                "Gate C1 повторяет только temporal часть frozen benchmark: четыре обязательные compact architectures × 11 outer folds × 5 seeds. В каждом outer train используются ровно три forward-only inner folds; 9 240 logical inner evaluations сведены к 3 640 физическим fits только при полном совпадении model/parameter/seed/sample/sequence/target/preprocessing/code/environment hashes. После tuning выполнено 220 outer refits. Spatial и transition designs намеренно не открываются до C2 admission.",
            ]
        )
        self.figure(FIGURES["b6_temporal"], "Temporal screen MAE моделей Gate B6")
        self.heading("6.7.2 Temporal качество", 3)
        top = self.b6a["focus_temporal_metrics"]
        self.table(
            "Ключевые rolling-origin metrics",
            ["Model", "MAE", "Median AE", "RMSE", "P95 AE", "B1 skill"],
            [
                (
                    item["model_id"],
                    f'{item["mae"]:.3f}',
                    f'{item["median_absolute_error"]:.3f}',
                    f'{item["rmse"]:.3f}',
                    f'{item["p95_absolute_error"]:.3f}',
                    f'{100 * item["b1_skill"]:.1f}%',
                )
                for item in top
            ],
            [4.2, 2.2, 2.5, 2.2, 2.2, 2.2],
        )
        self.paragraphs(
            [
                f'B7 показал минимальную rolling MAE {self.b6a["b7_rolling_mae_mm_per_year"]:.3f} мм/год против {self.b6a["b1_rolling_mae_mm_per_year"]:.3f} у B1, то есть skill {100 * self.b6a["b7_skill_vs_b1"]:.1f}%. B8 занял второе место с MAE 5,748 мм/год. Лучшие новые tabular models остались слабее B7 по headline metric.',
                "Средняя ошибка не описывает хвост: у B7 RMSE близка к 9,98 мм/год, P95 absolute error — 17,54 мм/год, а max error превышает 105 мм/год. Следовательно, редкие origins остаются критичными; оптимизация только средней MAE недостаточна.",
            ]
        )
        self.figure(FIGURES["b6_rolling"], "MAE по 11 rolling target dates")
        c1_status = self.c1_screen.set_index("model_id")["status"].to_dict()
        c1_status_display = {
            "PASSED_TEMPORAL_SCREEN": "PASSED",
            "REJECTED_TEMPORAL_SCREEN": "REJECTED",
            "REJECTED_MODEL_EXECUTION": "EXECUTION REJECT",
        }
        self.table(
            "Temporal metrics Gate C1 на одинаковых 595 origins",
            ["Model", "Role/status", "MAE", "Median fold MAE", "RMSE", "P95 AE", "B1 skill"],
            [
                (
                    row.model_id.split("_")[0],
                    "Context"
                    if str(row.model_id).startswith("B")
                    else c1_status_display[c1_status[row.model_id]],
                    f"{row.mae:.3f}",
                    f"{row.median_fold_mae:.3f}",
                    f"{row.rmse:.3f}",
                    f"{row.p95_absolute_error:.3f}",
                    f"{100 * row.b1_skill:.1f}%",
                )
                for row in self.c1_canonical.itertuples(index=False)
            ],
            [1.5, 4.1, 2.0, 2.8, 2.0, 2.1, 2.0],
        )
        c1_best = self.c1_canonical.loc[
            self.c1_canonical["model_id"].astype(str).str.startswith("C0")
        ].sort_values(["mae", "model_id"], kind="mergesort").iloc[0]
        b1_c1 = self.c1_canonical.loc[
            self.c1_canonical["model_id"].eq("B1_persistence_last_rate")
        ].iloc[0]
        b7_c1 = self.c1_canonical.loc[
            self.c1_canonical["model_id"].eq("B7_two_regime_imm")
        ].iloc[0]
        admitted_text = ", ".join(self.c1_admission["admitted_model_ids"]) or "нет"
        self.paragraphs(
            [
                f'Лучшая deep-архитектура по canonical mean-of-five-seeds — {c1_best["model_id"]}: MAE {c1_best["mae"]:.3f} мм/год. На том же universe B1 имеет {b1_c1["mae"]:.3f}, действующий primary B7 — {b7_c1["mae"]:.3f} мм/год. Это temporal train-only comparison, а не финальный рейтинг.',
                f'Программный C1 admission сформирован без weighted score: pooled и median fold MAE должны быть не хуже B1 более чем на 10%, worst fold — не более 2× B1, а execution/leakage checks обязаны пройти полностью. Допущенные к C2 deep model IDs: {admitted_text}. Низкое качество получает REJECTED_TEMPORAL_SCREEN и не считается software failure.',
            ]
        )
        self.figure(FIGURES["c1_temporal"], "Gate C1: canonical temporal MAE относительно B1, B7 и B8")
        self.figure(FIGURES["c1_rolling"], "Gate C1: MAE по 11 rolling target dates")
        self.heading("6.7.3 Spatial устойчивость", 3, page_break_before=True)
        spatial_focus = [
            item
            for item in self.b6a["focus_spatial_metrics"]
            if item["model_id"] in {"B1_persistence_last_rate", "B7_two_regime_imm", "B8_student_t_robust_imm"}
        ]
        self.table(
            "Spatial macro и worst-group MAE",
            ["Model", "Design", "Scope", "MAE, мм/год"],
            [
                (
                    item["model_id"].split("_")[0],
                    "Profile" if "profile" in item["design"] else "Zone",
                    {
                        "equal_profile_macro": "Equal-profile macro",
                        "worst_profile": "Worst profile",
                        "equal_zone_macro": "Equal-zone macro",
                        "worst_zone": "Worst zone",
                    }[item["scope"]],
                    f'{item["mae"]:.3f}',
                )
                for item in spatial_focus
            ],
            [2.4, 2.6, 7.2, 3.8],
        )
        self.paragraphs(
            [
                "B7 улучшает equal-profile macro MAE с 6,592 у B1 до 5,676 мм/год и worst-profile с 22,238 до 17,809 мм/год. Equal-zone macro улучшается с 5,883 до 4,975, worst-zone — с 9,563 до 8,478 мм/год. B8 показывает немного лучшую equal-zone macro MAE 4,928, но худший worst-zone 8,920 и более слабую rolling MAE.",
                "Новые tabular models теряют устойчивость в leave-zone design. Например, Z01 ElasticNet имеет equal-zone macro 8,344 и worst-zone 12,948 мм/год, а XGBoost — 7,807 и 15,172. Именно spatial guards не позволили заменить B7 моделью, которая выглядит приемлемо только на pooled temporal evidence.",
            ]
        )
        self.figure(FIGURES["b6_spatial"], "Spatial stability frozen comparators и shortlisted models")
        self.heading("6.7.4 Transition-specific качество", 3)
        transitions = [
            item
            for item in self.b6a["focus_transition_metrics"]
            if item["model_id"] in {"B1_persistence_last_rate", "B7_two_regime_imm", "B8_student_t_robust_imm"}
            and item["dimension"] in {"pooled_transition", "transition"}
        ]
        self.table(
            "Transition metrics frozen comparators",
            ["Model", "Сегмент", "Origins", "Профили", "MAE, мм/год"],
            [
                (
                    item["model_id"].split("_")[0],
                    item["segment"],
                    item["rows"],
                    item["profiles"],
                    f'{item["mae"]:.3f}',
                )
                for item in transitions
            ],
            [2.2, 5.2, 2.4, 2.4, 3.8],
        )
        self.paragraphs(
            [
                "На pooled transition B7 снижает MAE с 9,990 у B1 до 9,169 мм/год, но это улучшение около 8,2% и не достигает preregistered 10%. На accelerating B7 заметно лучше B1: 11,491 против 13,392. На volatile-or-gap различие почти отсутствует: 7,005 против 7,027.",
                "B8 не исправляет узкую проблему: accelerating MAE повышается до 12,748, volatile-or-gap — до 7,068, pooled transition — до 9,572. Результат подтверждает, что robust observation likelihood не эквивалентна лучшей transition dynamics.",
            ]
        )
        self.figure(FIGURES["b6_transition"], "Ошибка моделей по transition segments")
        self.heading("6.7.5 Интервальная калибровка", 3)
        conformal = self.b6a["focus_conformal_metrics"]
        self.table(
            "Scaled conformal metrics",
            ["Model", "Coverage 50", "Coverage 80", "Coverage 95", "Width 95", "WIS"],
            [
                (
                    item["model_id"].split("_")[0],
                    f'{item["coverage_50"]:.3f}',
                    f'{item["coverage_80"]:.3f}',
                    f'{item["coverage_95"]:.3f}',
                    f'{item["mean_width_95"]:.2f}',
                    f'{item["weighted_interval_score"]:.3f}',
                )
                for item in conformal
            ],
            [2.3, 2.7, 2.7, 2.7, 2.8, 2.3],
        )
        self.paragraphs(
            [
                "B7 достигает coverage 0,951 для 95% интервала при средней ширине 51,40 мм/год и WIS 3,788 — лучшем среди показанных focus models. B1 имеет сходное coverage 0,950, но ширину 56,43 и WIS 4,171. Это означает, что B7 одновременно сохраняет calibration и повышает sharpness.",
                "Coverage около nominal не снимает вопрос conditional calibration. Профили, зоны и volatile/gap сегменты имеют меньше наблюдений; conditional coverage публикуется только при support не менее 30 origins. Нельзя переносить pooled coverage на каждый профиль как доказанный факт.",
            ]
        )
        self.figure(FIGURES["b6_calibration"], "Empirical coverage и ширина conformal intervals")
        self.heading("6.7.6 Learning curves и неопределённость сравнений", 3)
        self.paragraphs(
            [
                "Learning curves строятся на неизменном audit tail; увеличение числа кампаний меняет не только объём, но и историческое покрытие regimes. Параметры на кривых не перенастраиваются, поэтому они отражают data sufficiency, а не новый tuning loop.",
                "Paired absolute-error deltas оцениваются двумя sensitivity procedures по 2 000 replicates: resampling 14 profile clusters и 11 target-date blocks с seed 42117. Наивный i.i.d. bootstrap по origins не используется. Из-за четырёх зон inferential confidence interval для zone effect не интерпретируется; публикуются все значения и worst-zone [15].",
                "Для Gate C1 отдельно проверена seed stability. Все пять заранее фиксированных seeds участвуют и в tuning, и в outer refit; их нельзя отбрасывать по качеству. Пороговые признаки IQR ≤ 0,50 мм/год и CV ≤ 10% публикуются сейчас описательно, но становятся обязательными только при будущей suite-v5 eligibility в Gate C2.",
            ]
        )
        self.table(
            "Пятиseedовая устойчивость compact sequence-моделей",
            ["Model", "Mean seed MAE", "IQR", "CV", "Range", "Ensemble MAE", "Дат лучше B1/B7"],
            [
                (
                    row.model_id.split("_")[0],
                    f"{row.seed_mae_mean:.3f}",
                    f"{row.seed_mae_iqr:.3f}",
                    f"{100 * row.seed_mae_cv:.2f}%",
                    f"{row.seed_mae_range:.3f}",
                    f"{row.ensemble_mae:.3f}",
                    f"{int(row.dates_improved_vs_b1)}/{int(row.dates_improved_vs_b7)}",
                )
                for row in self.c1_seeds.itertuples(index=False)
            ],
            [1.6, 2.8, 1.8, 1.8, 1.9, 2.7, 3.4],
        )
        self.figure(FIGURES["b6_learning"], "Learning curves frozen comparators и shortlisted models")
        self.figure(FIGURES["b6_sensitivity"], "Profile-cluster sensitivity paired deltas относительно B7")
        self.figure(FIGURES["c1_seed"], "Gate C1: разброс пяти fixed-seed MAE и canonical ensemble")
        self.heading("6.7.7 Итог Gate B6 и переход к C", 3)
        self.status_box(
            f'Gate B6: {self.b6["status"]}. Screened models: {self.b6["models_screened"]}; advanced: {self.b6["models_advanced"]}; new eligible: {self.b6["new_models_eligible"]}; primary: {self.b6["primary_model_id"]}. Final quality claim allowed: {str(self.b6["final_quality_claim_allowed"]).lower()}.'
        )
        self.paragraphs(
            [
                "Suite v4 заморожена до нового holdout. B1/B5/B6/B7/B8 сохраняются как context comparators, а лучший interpretable/probabilistic result не может post-hoc изменить primary после holdout. Такое решение предотвращает переобучение исследовательского процесса, даже если отдельная новая модель выигрывает на одном сегменте.",
                f'Gate C1 проверил новый класс hypotheses без продолжения тюнинга B7 и завершился статусом {self.c1_validation["status"]}. Он сформировал только temporal admission manifest; profile/zone/transition/calibration audit и suite v5 остаются PENDING. Поэтому suite v4 и B7 primary в настоящей редакции не изменены.',
            ]
        )

    def _section_68(self) -> None:
        self.heading("6.8 Error Atlas и ограничения мониторинга", 2)
        self.heading("6.8.1 Структура Error Atlas", 3)
        self.paragraphs(
            [
                "Error Atlas агрегирует абсолютную ошибку по target date, profile, zone, point, transition, horizon, n_history, uncertainty и missing-campaign bins. Одновременно публикуются pooled micro, equal-profile/equal-zone macro, worst group и worst 10% points. Каждая строка содержит support status, поэтому низкоподдержанные сегменты не влияют на выбор.",
                "Residual dependence оценивается descriptively внутри профиля и календарной даты; отдельно фиксируется число independent temporal/profile/zone units. Это не превращает 911 origins в 911 независимых экспериментов. Cluster resampling и leave-one-profile-out jackknife показывают, насколько вывод зависит от отдельных профилей.",
            ]
        )
        self.heading("6.8.2 Основные ограничения", 3)
        self.table(
            "Ограничения и меры контроля",
            ["Ограничение", "Риск", "Принятая мера"],
            [
                ("Реконструированные/синтетические данные", "Нет прямой промышленной валидности", "Claim boundary; future/external holdout"),
                ("98 trajectories и 14 profiles", "Псевдорепликация", "Temporal + profile/zone folds; cluster sensitivity"),
                ("3–16 наблюдений", "Высокая variance сложных моделей", "Compact parameter budget; B7 fallback"),
                ("42–560 дней между эпохами", "Неверная regular-grid динамика", "Explicit delta_t и missing masks"),
                ("Proxy zones", "Не инженерное районирование", "Split-only use; worst-zone descriptive"),
                ("58 opaque sample IDs", "Ошибочная семантическая интерпретация", "IDs metadata-only; future versioning"),
                ("Старый test раскрыт", "Selection contamination", "Не использовать; новый policy/intake"),
                ("Нового holdout нет", "Нет финальной оценки", "PENDING_DATA; one-access rule"),
            ],
            [4.2, 5.2, 7.1],
        )
        self.heading("6.8.3 Ограничения Gate C", 3)
        self.paragraphs(
            [
                "Sequence models видят много overlapping prefixes одной траектории. Это увеличивает число training origins, но не число независимых физических объектов. Поэтому multi-seed stability и spatial holdouts обязательны, а parameter budget ограничен 100 000. Большая transformer-модель не считается более научной только из-за архитектурной новизны.",
                "Left padding до 16 не создаёт искусственных наблюдений: padding rows маскируются и не участвуют в preprocessing fit. Если будущая история станет длиннее 16, протокол заранее предписывает удалять самые старые tokens. Изменение длины контекста после появления holdout потребует новой версии suite и policy.",
                "Current campaign type и target campaign type являются планово известными категориями, но campaign IDs запрещены. Это разделяет доступный календарный контекст и memorization конкретной даты. Аналогично point/profile/zone используются для resampling и диагностики, но не как embeddings.",
                "Gate C1 ограничен temporal screen. Даже PASSED_TEMPORAL_SCREEN не означает spatial или transition eligibility: leave-profile-out, leave-zone-out, transition/gap audit и common conformal calibration не вычислялись. Пять seeds характеризуют algorithmic variability, но не заменяют независимые физические объекты и новый внешний holdout.",
            ]
        )
        self.heading("6.8.4 Требования к эксплуатационному продолжению", 3, page_break_before=True)
        self.numbered(
            [
                "Получить новый пакет минимум из 100 observed origins, 75 точек, 12 профилей и двух target campaign dates либо внешний независимый объект с эквивалентным target definition.",
                "До чтения targets заморозить origins manifest, schema mapping, hashes, candidate commit и suite v5 primary.",
                "Создать новую holdout policy/intake version; не переписывать v3.",
                "Выполнить ровно одну evaluation frozen primary; failed access считать использованным.",
                "После доступа запретить tuning и model selection; результаты comparators публиковать только context-only.",
                "Провести инженерную интерпретацию ошибок совместно со специалистом предметной области и заменить proxy zones на утверждённое районирование при его наличии.",
            ]
        )

    def _conclusion(self) -> None:
        self.heading("6.9 Выводы по специальной части", 2)
        self.paragraphs(
            [
                "Построен воспроизводимый контур прогнозирования next-planned скорости оседания T1. Канонические таблицы и contracts отделены от historical/private/evaluation-only данных; split manifests и model-facing loaders предотвращают случайное открытие test и random row resampling. Grain 1 274 origins → 98 trajectories → 14 profiles учтён temporal, rolling, profile и zone designs.",
                "На расширенном train-only benchmark лучшим frozen candidate остаётся B7 two-regime IMM: rolling MAE 5,640 мм/год, skill 10,6% относительно B1, equal-profile macro 5,676 и equal-zone macro 4,975 мм/год. B7 сохраняет 95% conformal coverage 0,951 при WIS 3,788. При этом pooled transition improvement не достигает preregistered 10%, а rare extreme errors сохраняются; поэтому результат не интерпретируется как окончательный.",
                "Gate B6 корректно завершился PASS_NO_NEW_PRIMARY. Новые tabular/probabilistic models не прошли одновременно temporal, transition, spatial и calibration criteria. Это подтверждает необходимость governance, в котором отдельная победа на pooled MAE не перекрывает worst-zone или sign-consistency failure.",
                f'Gate C0 заморозил sequence representation всех 911 train origins, feature/mask contract и causal fold bindings. Gate C1 затем выполнил четыре обязательные compact sequence architectures на 11 rolling folds и пяти seeds, сохранил 11 900 single-seed и 2 380 ensemble deep prediction rows, а также 1 785 frozen comparator rows. Независимый validator завершился статусом {self.c1_validation["status"]}.',
                f'К Gate C2 допущены: {", ".join(self.c1_admission["admitted_model_ids"]) or "ни одна deep-модель"}. Это только temporal admission. Следующий этап должен выполнить profile/zone/transition audit и calibration без исторического validation/test; suite v5 пока не создана, suite v4 с B7 не изменена. Окончательный вывод возможен только после нового real future/external holdout; статус финальной оценки — PENDING_DATA.',
            ]
        )

    def _references(self) -> None:
        self.heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 2)
        self.paragraph(
            "Сведения о нормативных документах сверены с официальным порталом Росстандарта по состоянию на 01.09.2026. Локальные machine artifacts перечислены с repository-relative путями для воспроизводимости.",
        )
        for index, reference in enumerate(REFERENCES, start=1):
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Cm(0.75)
            paragraph.paragraph_format.first_line_indent = Cm(-0.75)
            paragraph.paragraph_format.keep_together = True
            paragraph.add_run(f"{index}. {reference}")

    def save(self) -> None:
        OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
        work = ROOT / "work" / "gate_c1_reporting"
        work.mkdir(parents=True, exist_ok=True)
        temporary_docx = work / f"{OUTPUT_PATH.stem}.{uuid4().hex}.docx"
        self.doc.save(temporary_docx)
        temporary_docx.replace(OUTPUT_PATH)
        source_rows = []
        for name, path in {**DATA_PATHS, **FIGURES}.items():
            source_rows.append(
                {
                    "name": name,
                    "path": path.relative_to(ROOT).as_posix(),
                    "bytes": path.stat().st_size,
                    "sha256": _sha256_file(path),
                }
            )
        source_map = {
            "schema_version": 1,
            "document": OUTPUT_PATH.relative_to(ROOT).as_posix(),
            "document_bytes": OUTPUT_PATH.stat().st_size,
            "document_sha256": _sha256_file(OUTPUT_PATH),
            "generated_on": "2026-09-05",
            "design_preset": "gost_vkr_ru",
            "format_basis": ["GOST 7.32-2017", "GOST R 7.0.100-2018", "GOST R 7.0.5-2008"],
            "institutional_rules_override": True,
            "claim_boundary": "train_only_internal_research_no_final_quality_claim",
            "gate_c0_model_training_calls": 0,
            "gate_c1_logical_inner_evaluations": int(self.c1_compute["logical_inner_evaluations"].sum()),
            "gate_c1_physical_inner_fits": int(self.c1_compute["physical_inner_fits_executed"].sum()),
            "gate_c1_outer_refits": int(self.c1_compute["outer_refits"].sum()),
            "gate_c1_model_training_calls": int(
                self.c1_compute["physical_inner_fits_executed"].sum()
                + self.c1_compute["outer_refits"].sum()
            ),
            "gate_c1_reporting_training_calls": 0,
            "historical_validation_loaded_for_gate_c1": False,
            "current_test_loaded_for_gate_c1": False,
            "new_holdout_seen_for_gate_c1": False,
            "profile_zone_transition_audit_executed": False,
            "suite_v5_created": False,
            "numbered_table_count": self.table_number,
            "physical_table_count": len(self.doc.tables),
            "figure_count": self.figure_number,
            "sources": source_rows,
        }
        temporary_map = work / f"{SOURCE_MAP_PATH.name}.{uuid4().hex}.tmp"
        temporary_map.write_text(
            json.dumps(source_map, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
            newline="\n",
        )
        temporary_map.replace(SOURCE_MAP_PATH)
        print(json.dumps(source_map, ensure_ascii=False, indent=2))


def _style_heading(style, *, size: int, alignment, page_break: bool, before: int, after: int) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.italic = False
    style.font.color.rgb = RGBColor(0, 0, 0)
    _set_east_asia_font(style, "Times New Roman")
    style.paragraph_format.alignment = alignment
    style.paragraph_format.first_line_indent = Cm(0)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
    style.paragraph_format.keep_together = True
    style.paragraph_format.page_break_before = page_break


def _set_east_asia_font(style, font_name: str) -> None:
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _set_cell_text(cell, text: str, document: Document, *, alignment) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles["GOST Table Text"]
    paragraph.alignment = alignment
    paragraph.paragraph_format.keep_together = True
    paragraph.add_run(text)
    _set_cell_margins(cell, top=70, start=90, bottom=70, end=90)


def _set_cell_margins(cell, *, top: int, start: int, bottom: int, end: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_table_fixed_layout(table) -> None:
    table_pr = table._tbl.tblPr
    layout = table_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        table_pr.append(layout)
    layout.set(qn("w:type"), "fixed")


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    tr_pr.append(table_header)


def _prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cannot_split = OxmlElement("w:cantSplit")
    tr_pr.append(cannot_split)


def _add_field(run, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, result, end])


def _sha256_file(path: Path) -> str:
    digest = sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def main() -> int:
    builder = SpecialSectionBuilder()
    builder.build()
    builder.save()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
